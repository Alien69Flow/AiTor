from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "AiTor_Modulo_1_Plano_de_Entrega.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5E6873"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF4EE"
PALE_GOLD = "FFF4D6"
PALE_RED = "FBEAEC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_run(run, size=11, color="000000", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def add_text(doc, text, style="Normal", size=None, color=None, bold=False, italic=False, after=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run(r, size or 11, color or "000000", bold, italic)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run(r)
    return p


def write_cell(cell, text, bold=False, color="000000", size=9.4, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=bold)


def add_matrix(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        write_cell(cell, text, bold=True, color=INK, size=9.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, row_data)):
            fill = None
            if status_col is not None and index == status_col:
                lowered = text.lower()
                if "conclu" in lowered or "dispon" in lowered:
                    fill = PALE_GREEN
                elif "andamento" in lowered or "validar" in lowered:
                    fill = PALE_GOLD
                elif "pendente" in lowered or "não" in lowered:
                    fill = PALE_RED
            if fill:
                set_cell_shading(cell, fill)
            write_cell(cell, text, size=9.1, align=WD_ALIGN_PARAGRAPH.CENTER if index == status_col else WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label} ")
    set_run(r, 10.5, DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run(r, 10.5, "000000")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run(r, {1: 16, 2: 13, 3: 12}[level], {1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, before, after, size, color in ((1, 18, 10, 16, BLUE), (2, 14, 7, 13, BLUE), (3, 10, 5, 12, DARK_BLUE)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    run = paragraph.add_run("Página ")
    set_run(run, 8.5, MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header_p = section.header.paragraphs[0]
    clear_paragraph(header_p)
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header_p.add_run("AiTor v69 | Plano de entrega do Módulo 1")
    set_run(r, 8.5, MUTED, bold=True)
    footer_p = section.footer.paragraphs[0]
    clear_paragraph(footer_p)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer_p)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ALIENFLOW / AITOR")
    set_run(r, 10, DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Módulo 1 - Plano de conclusão e entrega")
    set_run(r, 23, INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Escopo consolidado para o MVP de monetização web + Telegram")
    set_run(r, 13.5, MUTED)

    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [1800, 7560])
    for label, value in (("Base", "Conversa João/Aitor + revisão técnica do repositório atual"), ("Situação", "Módulo 1 em evolução; ainda não é uma entrega final de paywall"), ("Data de referência", "15 de agosto de 2026")):
        row = meta.rows[["Base", "Situação", "Data de referência"].index(label)]
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        write_cell(row.cells[0], label, bold=True, color=INK, size=9.5)
        write_cell(row.cells[1], value, size=9.5)
    doc.add_paragraph()

    add_callout(doc, "Leitura executiva.", "A base do Módulo 1 já existe: bot/webhooks, miniapp/web, agentes, RAG, loops e um paywall visual. A entrega ainda depende de validação ponta a ponta, estabilidade do Globe e de uma integração real de wallet + pagamento/entitlement em crypto ou NFT.")

    add_heading(doc, "1. O que esta primeira entrega precisa provar")
    add_text(doc, "Entregar um MVP utilizável no qual a pessoa entra pela web ou pelo Telegram, conecta a wallet, tem o acesso pago ou por NFT verificado no backend e recebe o mesmo nível de acesso em ambos os canais. O Globe e os fluxos de agentes devem permanecer utilizáveis durante esse percurso.")
    add_bullet(doc, "Uma fonte única de verdade para acesso: perfil, wallet, pagamento/NFT e tier do usuário.")
    add_bullet(doc, "Acesso consistente entre AiTor web/miniapp e o bot do Telegram.")
    add_bullet(doc, "Paywall real: não apenas tela de planos ou alerta visual.")
    add_bullet(doc, "Camadas do Globe combináveis sem desaparecer ou gerar falha bloqueante.")

    add_heading(doc, "2. Onde estamos agora")
    add_matrix(doc,
        ["Frente", "O que já existe", "Situação para entrega"],
        [
            ("Telegram, webhooks e respostas", "Bot, endpoints e callbacks no código; link do bot e versões web foram fornecidos.", "Base disponível; validar ponta a ponta"),
            ("Web / miniapp", "Interfaces e integrações Supabase existentes; versões implantadas indicadas por Aitor.", "Base disponível; auditar fluxo atual"),
            ("Agentes, RAG e loops", "Supervisor, agentes especializados, bases RAG e loops já estruturados.", "Em andamento; definir quais loops entram no MVP"),
            ("Globe", "Controles de camadas e Cesium existentes; correção recente para filtros e arcos por camada.", "Em andamento; regressão de combinações pendente"),
            ("Wallet com Reown", "Objetivo informado por Aitor; não há integração Reown funcional no repositório revisado.", "Pendente"),
            ("Paywall crypto / NFT", "Planos e créditos visuais existem; checkout, confirmação on-chain e entitlement não.", "Pendente"),
            ("Repositório e rastreabilidade", "Código atualizado publicado no repositório privado kyzencodex/aitor com log de implementação.", "Concluído")
        ], [2100, 4700, 2560], status_col=2)
    add_text(doc, "Conclusão: o Módulo 1 não deve ser tratado como “feito” ainda. Ele tem uma fundação relevante, mas o fluxo que monetiza e libera acesso precisa ser fechado e testado.", italic=True, color=DARK_BLUE, after=8)

    add_heading(doc, "3. Escopo proposto do Módulo 1", 1)
    add_heading(doc, "Dentro da entrega", 2)
    add_bullet(doc, "Validar o percurso atual web + Telegram + webhook com uma conta de teste.")
    add_bullet(doc, "Estabilizar e testar as combinações prioritárias das camadas de dados do Globe.")
    add_bullet(doc, "Conectar wallet com Reown, relacionando-a ao usuário de forma segura.")
    add_bullet(doc, "Implementar verificação de pagamento em crypto e/ou posse de NFT no backend; atualizar o tier de acesso somente após verificação.")
    add_bullet(doc, "Aplicar a mesma regra de acesso na web/miniapp e no Telegram.")
    add_bullet(doc, "Registrar configurações necessárias, testes feitos, limitações conhecidas e procedimento de deploy.")
    add_heading(doc, "Fora desta primeira entrega", 2)
    add_bullet(doc, "Autoposting completo em redes sociais, CRM, A/B testing e analytics avançado.")
    add_bullet(doc, "Automação autônoma de trading ou publicação sem aprovação humana.")
    add_bullet(doc, "Expansão de novas features do Globe sem primeiro estabilizar o conjunto atual.")

    add_heading(doc, "4. Frentes de trabalho e critério de pronto")
    add_matrix(doc,
        ["Frente", "Trabalho agora", "Pronto quando"],
        [
            ("A. Integração base", "Confirmar updates do bot, webhook, miniapp e mapeamento de usuário entre canais.", "Um usuário de teste completa o fluxo sem quebra de contexto."),
            ("B. Globe", "Executar matriz de combinações, registrar falhas e corrigir filtros/dados que somem ao combinar camadas.", "Combinações prioritárias aparecem; nenhum erro bloqueante no console."),
            ("C. Loops e agentes", "Escolher loops do MVP, colocar estado/logs mínimos e manter ações sensíveis em modo aprovado/manual.", "O operador vê o que rodou, seu resultado e pode interromper o loop."),
            ("D. Wallet", "Configurar Reown, redes aceitas, conexão, assinatura de prova de posse e vínculo com perfil.", "Wallet conectada e vinculada de forma verificável ao usuário."),
            ("E. Crypto / NFT", "Criar pedido, verificar transação ou posse de NFT no backend, conceder/renovar tier e tratar falhas.", "Acesso só é liberado após confirmação; recarga não duplica crédito."),
            ("F. Release", "Build, variáveis de ambiente, smoke test web/Telegram, changelog e deploy.", "Fluxo demonstrável e documentação entregue."),
        ], [1700, 4050, 3610])

    add_heading(doc, "5. Sequência recomendada para fechar o MVP")
    add_number(doc, "Congelar o baseline: confirmar que o repositório e as versões implantadas apontam para a mesma versão. Esta organização já foi iniciada.")
    add_number(doc, "Validar primeiro o que existe: bot, webhook, miniapp, acesso atual e combinações do Globe. Corrigir somente regressões que bloqueiem esse percurso.")
    add_number(doc, "Fechar a regra de negócio Web3: rede, ativo, preço, recorrência, carteira de recebimento e regra exata de NFT/tier.")
    add_number(doc, "Implementar wallet e entitlement: conexão Reown, prova de posse, verificação no backend e persistência do tier.")
    add_number(doc, "Ligar o paywall real nos dois canais: compra/verificação na web e permissão/upsell equivalente no Telegram.")
    add_number(doc, "Fazer release controlado: build, smoke test, evidência do fluxo e nota de versão.")

    add_heading(doc, "6. Decisões que precisam vir de Aitor")
    add_callout(doc, "Sem essas definições,", "é possível preparar a estrutura, mas não é seguro finalizar cobrança ou liberar acessos de produção.")
    add_matrix(doc,
        ["Decisão", "Definição necessária"],
        [
            ("Rede e Reown", "Chains suportadas e Project ID de Reown (armazenado como variável de ambiente, nunca no código)."),
            ("Modelo de cobrança", "Moeda/token aceito, preço, recorrência ou pagamento único e carteira/serviço recebedor."),
            ("NFT", "Contrato, chain, coleção/token IDs e regra: posse concede qual tier e por quanto tempo?"),
            ("Entitlement", "Quais tiers existem, benefícios, créditos e comportamento na expiração/revogação."),
            ("Fluxo Telegram", "Como a pessoa parte do bot para conectar/pagar e como retorna ao bot com acesso liberado."),
            ("Produção", "URLs oficiais, responsável pelo deploy e acesso seguro às variáveis de ambiente."),
        ], [2050, 7310])

    add_heading(doc, "7. Responsabilidades propostas")
    add_matrix(doc,
        ["Responsável", "Responsabilidade"],
        [
            ("Aitor / produto", "Definir regra comercial e Web3, fornecer acessos/variáveis por canal seguro, aprovar experiência e validar o fluxo real."),
            ("João / implementação", "Auditar baseline, executar correções do Globe, integrar wallet/paywall após decisões, validar web + Telegram e publicar documentação técnica."),
            ("Ambos", "Aceitar o critério de pronto, testar uma conta real/de teste e aprovar a entrega antes de ampliar o escopo."),
        ], [2050, 7310])

    add_heading(doc, "8. Critério final de aceitação")
    add_bullet(doc, "A build de produção conclui e as variáveis exigidas estão documentadas sem expor segredos.")
    add_bullet(doc, "Uma wallet é conectada e comprovada; o vínculo persiste para o usuário correto.")
    add_bullet(doc, "Pagamento crypto ou NFT é confirmado no backend antes da concessão de tier/créditos.")
    add_bullet(doc, "Web/miniapp e Telegram reconhecem o mesmo direito de acesso e exibem uma ação clara quando não há acesso.")
    add_bullet(doc, "As combinações de camadas prioritárias do Globe foram testadas e não escondem dados indevidamente.")
    add_bullet(doc, "Loops incluídos no MVP têm estado/log mínimo e não executam ação irreversível sem regra aprovada.")
    add_bullet(doc, "Changelog, limitações conhecidas e roteiro de demonstração foram entregues junto ao código.")

    add_heading(doc, "9. Estado atual e próximo passo")
    add_callout(doc, "Fase atual: estabilização e validação do baseline.", "O repositório foi organizado, o controle de camadas do Globe recebeu uma correção focada e a build de produção foi validada. O próximo passo prático é executar a validação ponta a ponta do fluxo atual e fechar as decisões de rede, cobrança e NFT para iniciar a implementação real do entitlement.")

    doc.core_properties.title = "AiTor v69 - Módulo 1: Plano de conclusão e entrega"
    doc.core_properties.subject = "Escopo, estado atual e critérios de entrega do MVP"
    doc.core_properties.author = "AlienFlow / AiTor"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
